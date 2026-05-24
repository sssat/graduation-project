from __future__ import annotations

import json
import math
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
RESULTS_DIR = ROOT / "tools" / "load-test" / "results"
OUTPUT_DOCX = ROOT / "docs" / "newsight-load-test-report-with-causes.docx"
CHART_PATH = ROOT / "tools" / "load-test" / "results" / "load-test-chart.png"


@dataclass
class ResultRow:
    vus: int
    duration_seconds: int
    requests: int
    rps: float
    failure_rate: float
    avg_ms: float
    p90_ms: float
    p95_ms: float
    max_ms: float
    source: str
    note: str = ""


def metric_values(data: dict, name: str) -> dict:
    return data.get("metrics", {}).get(name, {}).get("values", {})


def load_result_rows() -> list[ResultRow]:
    rows: list[ResultRow] = [
        ResultRow(
            vus=5,
            duration_seconds=30,
            requests=380,
            rps=12.40,
            failure_rate=0.00,
            avg_ms=18.14,
            p90_ms=22.26,
            p95_ms=25.43,
            max_ms=47.12,
            source="terminal",
            note="사전 점검",
        )
    ]

    for path in sorted(RESULTS_DIR.glob("k6-summary-*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        vus = int(metric_values(data, "vus_max").get("value", 0))
        requests = int(metric_values(data, "http_reqs").get("count", 0))
        rps = float(metric_values(data, "http_reqs").get("rate", 0))
        failed = metric_values(data, "http_req_failed")
        duration = metric_values(data, "http_req_duration")
        rows.append(
            ResultRow(
                vus=vus,
                duration_seconds=60 if vus == 10 else 120,
                requests=requests,
                rps=rps,
                failure_rate=float(failed.get("rate", 0)) * 100,
                avg_ms=float(duration.get("avg", 0)),
                p90_ms=float(duration.get("p(90)", 0)),
                p95_ms=float(duration.get("p(95)", 0)),
                max_ms=float(duration.get("max", 0)),
                source=path.name,
            )
        )

    return sorted(rows, key=lambda item: item.vus)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color: str = "D8DEE8", size: str = "6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(table, top: int = 100, start: int = 100, bottom: int = 100, end: int = 100) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tbl_cell_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text: str, bold: bool = False, color: str | None = None, size: int | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_body_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        add_run(p, text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(55, 65, 81)
    run.bold = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def build_chart(rows: list[ResultRow]) -> None:
    width, height = 1500, 680
    margin_left, margin_right = 120, 80
    margin_top, margin_bottom = 125, 95
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom

    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)

    font_path = Path("C:/Windows/Fonts/malgun.ttf")
    bold_path = Path("C:/Windows/Fonts/malgunbd.ttf")
    font = ImageFont.truetype(str(font_path), 28) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 23) if font_path.exists() else ImageFont.load_default()
    title_font = ImageFont.truetype(str(bold_path), 36) if bold_path.exists() else font

    draw.text((margin_left, 24), "동시 가상 사용자 수에 따른 응답 시간 변화", fill="#1F4E79", font=title_font)
    draw.text((margin_left, 70), "150 VU부터 지연이 증가하고 200 VU에서 포화 징후 관찰", fill="#4B5563", font=small)

    max_ms = max(row.p95_ms for row in rows)
    y_max = math.ceil(max_ms / 250) * 250
    if y_max < 1250:
        y_max = 1250

    for i in range(0, 6):
        value = y_max * i / 5
        y = margin_top + chart_h - (value / y_max) * chart_h
        draw.line((margin_left, y, width - margin_right, y), fill="#E5E7EB", width=2)
        draw.text((35, y - 15), f"{int(value)}ms", fill="#6B7280", font=small)

    draw.line((margin_left, margin_top, margin_left, margin_top + chart_h), fill="#94A3B8", width=3)
    draw.line((margin_left, margin_top + chart_h, width - margin_right, margin_top + chart_h), fill="#94A3B8", width=3)

    spacing = chart_w / len(rows)
    bar_w = min(110, spacing * 0.42)
    points = []

    for idx, row in enumerate(rows):
        center_x = margin_left + spacing * idx + spacing / 2
        bar_left = center_x - bar_w / 2
        bar_right = center_x + bar_w / 2
        p95_top = margin_top + chart_h - (row.p95_ms / y_max) * chart_h
        avg_top = margin_top + chart_h - (row.avg_ms / y_max) * chart_h
        color = "#2F80ED" if row.vus < 200 else "#D97706"
        draw.rounded_rectangle((bar_left, p95_top, bar_right, margin_top + chart_h), radius=10, fill=color)
        points.append((center_x, avg_top))
        draw.text((center_x - 42, margin_top + chart_h + 18), f"{row.vus} VU", fill="#374151", font=small)
        draw.text((center_x - 52, p95_top - 34), f"{row.p95_ms:.0f}", fill="#111827", font=small)

    for first, second in zip(points, points[1:]):
        draw.line((first[0], first[1], second[0], second[1]), fill="#10B981", width=5)
    for x, y in points:
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill="#10B981")

    legend_y = height - 54
    draw.rounded_rectangle((margin_left, legend_y, margin_left + 32, legend_y + 20), radius=5, fill="#2F80ED")
    draw.text((margin_left + 42, legend_y - 5), "p95 응답 시간(ms)", fill="#374151", font=small)
    draw.line((margin_left + 310, legend_y + 10, margin_left + 370, legend_y + 10), fill="#10B981", width=5)
    draw.ellipse((margin_left + 335, legend_y + 2, margin_left + 351, legend_y + 18), fill="#10B981")
    draw.text((margin_left + 382, legend_y - 5), "평균 응답 시간(ms)", fill="#374151", font=small)

    image.save(CHART_PATH)


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Newsight 운영 서버 트래픽 부하 테스트 결과 보고서", color="6B7280", size=8)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Newsight", bold=True, color="1F4E79", size=18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    add_run(title, "운영 서버 트래픽 부하 테스트 결과 보고서", bold=True, color="111827", size=22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    add_run(subtitle, "k6 기반 공개 분석 조회 API 성능 검증", color="4B5563", size=12)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    set_table_borders(info, color="FFFFFF", size="0")
    set_cell_padding(info, top=110, start=170, bottom=110, end=170)
    values = [
        ("대상 시스템", "https://newsightkr.com/api"),
        ("테스트 도구", "Grafana k6"),
        ("테스트 일시", "2026-05-11"),
        ("테스트 범위", "공개 분석 조회 API(GET)"),
    ]
    for row, (label, value) in zip(info.rows, values):
        set_cell_shading(row.cells[0], "EAF2F8")
        set_cell_shading(row.cells[1], "F8FAFC")
        set_cell_text(row.cells[0], label, bold=True, color="1F4E79")
        set_cell_text(row.cells[1], value)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(box, color="D8DEE8")
    set_cell_padding(box, top=170, start=200, bottom=170, end=200)
    set_cell_shading(box.cell(0, 0), "F3F6FA")
    p = box.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        p,
        "요약 결론: 100 VU까지 안정, 150 VU부터 지연 증가, 200 VU에서 포화 징후 관찰",
        bold=True,
        color="1F4E79",
        size=11,
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_result_table(doc: Document, rows: list[ResultRow]) -> None:
    add_heading(doc, "4. 테스트 결과", 1)
    add_body_paragraph(
        doc,
        "모든 구간에서 HTTP 실패율은 0.00%로 측정되었다. 100 VU까지는 p95 응답 시간이 200ms 이하로 유지되었고, "
        "150 VU에서는 p95 응답 시간이 613.90ms로 증가하여 성능 저하가 시작되었다. 200 VU에서는 p95 응답 시간이 "
        "1,157.36ms로 증가하여 처리량 포화 징후가 관찰되었다.",
    )

    headers = ["구분", "시간", "요청 수", "RPS", "실패율", "평균", "p90", "p95", "최대"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_padding(table, top=95, start=90, bottom=95, end=90)

    widths = [1.55, 1.45, 1.8, 1.35, 1.45, 1.45, 1.45, 1.45, 1.55]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Cm(width)
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])

    for result in rows:
        cells = table.add_row().cells
        values = [
            f"{result.vus} VU" + (f"\n({result.note})" if result.note else ""),
            f"{result.duration_seconds}s",
            f"{result.requests:,}",
            f"{result.rps:.2f}",
            f"{result.failure_rate:.2f}%",
            f"{result.avg_ms:.2f}ms",
            f"{result.p90_ms:.2f}ms",
            f"{result.p95_ms:.2f}ms",
            f"{result.max_ms:.2f}ms",
        ]
        fill = "FFF7ED" if result.vus == 200 else "FFFFFF"
        for cell, value in zip(cells, values):
            set_cell_shading(cell, fill)
            set_cell_text(cell, value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    add_run(
        p,
        "주: 5 VU 결과는 저장 로직 수정 전 수행한 사전 점검 터미널 로그 기준이며, 10 VU 이상 결과는 저장된 k6 JSON/Markdown 결과 파일 기준이다.",
        color="6B7280",
        size=8,
    )


def add_endpoint_table(doc: Document) -> None:
    add_heading(doc, "3. 테스트 시나리오", 1)
    add_body_paragraph(
        doc,
        "테스트는 실제 운영 화면에서 자주 호출되는 공개 분석 조회 API를 가중치 기반으로 무작위 선택하여 반복 호출하는 방식으로 수행하였다. "
        "언론사 비교 분석은 하나의 통합 API가 아니라, 화면 구성 요소별 개별 API를 조합하여 조회하는 구조이므로 기사 수, 감성 비교, "
        "프레이밍 단어, 키워드 목록 API를 각각 테스트 대상에 포함하였다. 데이터 변경을 동반하는 회원가입, 문의 작성, 관리자 수정 API는 "
        "운영 데이터 보호를 위해 제외하였다.",
    )

    endpoints = [
        ("공통", "분석 개요", "/analytics/overview", "메인 분석 화면 상단 데이터 조회"),
        ("키워드 상세", "메타", "/analytics/keywords/{keyword_seq}", "키워드 기본 정보 조회"),
        ("키워드 상세", "AI 요약", "/analytics/keywords/{keyword_seq}/summary", "키워드 요약 조회"),
        ("키워드 상세", "워드클라우드", "/analytics/keywords/{keyword_seq}/wordcloud/title", "제목 주요 단어 조회"),
        ("키워드 상세", "검색 관심도", "/analytics/keywords/{keyword_seq}/search-timeline", "검색 관심도 시계열 조회"),
        ("키워드 상세", "본문 감성", "/analytics/keywords/{keyword_seq}/sentiment/content", "전체 본문 감성 조회"),
        ("키워드 상세", "제목 편향", "/analytics/keywords/{keyword_seq}/bias/title", "언론사별 제목 편향 조회"),
        ("언론사 비교", "키워드 목록", "/analytics/media-compare/keywords/top", "비교 대상 키워드 목록 조회"),
        ("언론사 비교", "기사 수", "/analytics/media-compare/keywords/{keyword_seq}/media-article-counts", "언론사별 기사 수 조회"),
        ("언론사 비교", "본문 감성", "/analytics/media-compare/keywords/{keyword_seq}/sentiment/content", "언론사별 본문 감성 비교"),
        ("언론사 비교", "프레이밍 단어", "/analytics/media-compare/keywords/{keyword_seq}/framing/title-top-words", "언론사별 제목 주요 단어 비교"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_padding(table, top=70, start=80, bottom=70, end=80)
    widths = [2.0, 2.2, 6.1, 4.1]
    for cell, header, width in zip(table.rows[0].cells, ["화면", "분류", "API 경로", "목적"], widths):
        cell.width = Cm(width)
        set_cell_shading(cell, "EAF2F8")
        set_cell_text(cell, header, bold=True, color="1F4E79")
    set_repeat_table_header(table.rows[0])

    for screen, category, endpoint, purpose in endpoints:
        cells = table.add_row().cells
        for cell, value, align_center in zip(cells, [screen, category, endpoint, purpose], [True, True, False, False]):
            set_cell_text(cell, value)
            if not align_center:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if value.startswith("/analytics"):
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(7.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    add_run(
        p,
        "주: 언론사 비교 화면은 키워드 목록 조회 후 선택된 keyword_seq를 기준으로 기사 수, 본문 감성, 프레이밍 단어 API를 병렬 호출한다. "
        "또한 화면 구성상 전체 본문 감성과 제목 편향 API도 함께 재사용된다.",
        color="6B7280",
        size=8,
    )


def add_summary_box(doc: Document) -> None:
    add_heading(doc, "1. 테스트 개요", 1)
    p = add_body_paragraph(doc)
    add_run(p, "본 문서는 Newsight 운영 서버의 공개 분석 조회 API를 대상으로 수행한 트래픽 부하 테스트 결과를 정리한 보고서이다. ")
    add_run(p, "테스트 목적", bold=True)
    add_run(p, "은 동시 접속 증가에 따른 응답 시간, 실패율, 처리량 변화를 확인하고 운영 환경의 안정 처리 가능 범위를 산정하는 것이다.")

    cards = [
        ("안정 처리 구간", "100 VU", "p95 190.65ms, 실패율 0.00%"),
        ("성능 저하 시작", "150 VU", "p95 613.90ms, 실패율 0.00%"),
        ("포화 관찰 구간", "200 VU", "p95 1,157.36ms, 실패율 0.00%"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CBD5E1")
    set_cell_padding(table, top=160, start=160, bottom=160, end=160)
    for cell, (title, value, detail) in zip(table.rows[0].cells, cards):
        set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=True, color="1F4E79", size=9)
        p.add_run("\n")
        add_run(p, value, bold=True, color="111827", size=16)
        p.add_run("\n")
        add_run(p, detail, color="4B5563", size=8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_environment(doc: Document) -> None:
    add_heading(doc, "2. 테스트 환경 및 기준", 1)
    rows = [
        ("테스트 대상", "https://newsightkr.com/api"),
        ("테스트 도구", "Grafana k6"),
        ("테스트 방식", "constant-vus 방식으로 지정한 가상 사용자 수를 테스트 시간 동안 유지"),
        ("가상 사용자 수", "10, 30, 50, 100, 150, 200 VU (5 VU 사전 점검 포함)"),
        ("테스트 시간", "10 VU: 60초, 30 VU 이상: 120초"),
        ("요청 간격", "각 VU가 응답 수신 후 0.25~0.5초 대기 후 다음 요청 수행"),
        ("성공 기준", "HTTP 실패율 1% 미만, p95 응답 시간 2초 미만"),
        ("테스트 데이터", "period=D7, keyword_seq 자동 선택"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_padding(table, top=105, start=120, bottom=105, end=120)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], "EAF2F8")
        set_cell_text(cells[0], label, bold=True, color="1F4E79")
        set_cell_text(cells[1], value)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_analysis(doc: Document, rows: list[ResultRow]) -> None:
    add_heading(doc, "5. 결과 분석", 1)
    doc.add_picture(str(CHART_PATH), width=Inches(6.05))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(caption, "그림 1. 동시 가상 사용자 수에 따른 평균/p95 응답 시간 변화", color="6B7280", size=8)

    add_heading(doc, "5.1 안정 구간", 2)
    add_body_paragraph(
        doc,
        "10, 30, 50 VU 구간에서는 요청 수가 증가하더라도 평균 응답 시간과 p95 응답 시간이 낮게 유지되었다. "
        "100 VU 구간에서도 실패율은 0.00%였으며 p95 응답 시간은 190.65ms로 측정되어, 공개 분석 조회 API 기준 안정적인 응답 성능을 보였다.",
    )
    doc.add_page_break()
    add_heading(doc, "5.2 성능 저하 시작 구간", 2)
    add_body_paragraph(
        doc,
        "150 VU 구간에서는 실패율이 0.00%로 유지되었으나 평균 응답 시간이 285.00ms, p95 응답 시간이 613.90ms로 증가하였다. "
        "100 VU 대비 요청 처리량 증가 폭은 제한적이었으므로, 이 구간부터 응답 지연이 본격적으로 증가하는 성능 저하 시작 구간으로 볼 수 있다.",
    )
    add_heading(doc, "5.3 포화 관찰 구간", 2)
    add_body_paragraph(
        doc,
        "200 VU 구간에서는 실패율이 여전히 0.00%였으나 평균 응답 시간이 533.08ms, p95 응답 시간이 1,157.36ms로 증가하였다. "
        "또한 150 VU 대비 요청 수 증가 폭이 거의 없고 응답 시간만 증가하여, 서버 또는 DB 조회 처리량이 포화되는 징후가 나타난 것으로 해석할 수 있다.",
    )


def add_suspected_causes(doc: Document) -> None:
    add_heading(doc, "5.4 병목 의심 원인 및 개선 가능성", 2)
    add_body_paragraph(
        doc,
        "100 VU 이후부터 요청 처리량(RPS)은 크게 증가하지 않는 반면 평균, p90, p95 응답시간이 급격히 증가하였다. "
        "이는 서버가 즉시 오류를 반환하는 상태라기보다, 애플리케이션 또는 DB 조회 처리 구간에서 대기 시간이 누적되는 형태로 해석된다. "
        "따라서 부하 테스트 결과만으로 단일 원인을 확정하기보다는, 백엔드 코드와 DB 조회 구조를 함께 확인하여 다음과 같은 개선 후보를 도출하였다.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_padding(table, top=90, start=90, bottom=90, end=90)

    headers = ["의심 원인", "근거", "개선 방향"]
    widths = [3.0, 5.4, 5.4]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Cm(width)
        set_cell_shading(cell, "EAF2F8")
        set_cell_text(cell, header, bold=True, color="1F4E79")
    set_repeat_table_header(table.rows[0])

    causes = [
        (
            "불필요한 전체 조회",
            "최신 분석 run 또는 랭킹 데이터를 가져올 때 필요한 1건/상위 N건보다 많은 데이터를 조회한 뒤 Java 코드에서 선택하는 구간이 확인되었다.",
            "DB 쿼리에서 LIMIT, Pageable, SUM 집계를 사용하여 애플리케이션으로 전달되는 데이터량을 줄인다.",
        ),
        (
            "DB 인덱스 활용 저하 가능성",
            "기사 발행일 조회에서 DATE(PUBLISHED_AT) 형태의 조건을 사용하고 있어 날짜 컬럼 인덱스를 충분히 활용하지 못할 가능성이 있다.",
            "PUBLISHED_AT >= 시작일 00:00:00 AND PUBLISHED_AT < 종료일 다음날 00:00:00 형태의 범위 조건으로 변경한다.",
        ),
        (
            "반복 조회 및 N+1 조회",
            "언론사별 워드클라우드 단어 조회처럼 목록을 먼저 조회한 뒤 반복문 안에서 세부 데이터를 다시 조회하는 구조가 일부 존재한다.",
            "여러 식별자를 IN 조건으로 한 번에 조회하고, 애플리케이션에서 그룹핑하여 DB 왕복 횟수를 줄인다.",
        ),
        (
            "DB 커넥션/서버 자원 한계",
            "150 VU와 200 VU에서 실패율은 0%이지만 RPS가 100 VU 대비 크게 늘지 않아, 커넥션 풀 또는 DB 처리량 포화 가능성이 있다.",
            "HikariCP 커넥션 풀, DB CPU/메모리, 느린 쿼리 로그를 확인하고 운영 서버 사양에 맞게 조정한다.",
        ),
        (
            "캐시 미적용",
            "분석 결과 API는 배치로 생성된 데이터를 조회하는 성격이 강해 동일 요청이 반복될 가능성이 높다.",
            "최신 run, 키워드 메타데이터, 조회 결과 일부에 짧은 TTL 캐시를 적용하여 반복 DB 조회를 줄인다.",
        ),
    ]

    for cause, evidence, direction in causes:
        cells = table.add_row().cells
        values = [cause, evidence, direction]
        for idx, (cell, value) in enumerate(zip(cells, values)):
            set_cell_shading(cell, "FFFFFF")
            set_cell_text(cell, value, bold=(idx == 0), color="111827")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT

    add_heading(doc, "5.5 개선 과정 기술 방향", 2)
    add_body_paragraph(
        doc,
        "보고서에는 개선 과정을 '부하 테스트 수행 -> 병목 징후 확인 -> 백엔드 코드 분석 -> 개선 후보 도출 -> 수정 후 재측정' 순서로 기술한다. "
        "특히 150 VU부터 응답시간이 증가했지만 실패율은 0%였다는 점을 근거로, 장애 대응보다는 조회 효율화와 DB 자원 사용량 감소를 목표로 개선을 진행한다고 설명하는 것이 적절하다.",
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "6. 결론 및 향후 개선 방향", 1)
    conclusions = [
        "운영 서버는 공개 분석 조회 API 기준으로 100 VU까지 안정적인 성능을 보였다.",
        "150 VU에서는 HTTP 실패 없이 정상 응답을 유지했지만, 100 VU 대비 응답 지연이 뚜렷하게 증가하였다.",
        "200 VU에서도 장애나 HTTP 실패는 발생하지 않았으나, 처리량 증가 폭이 제한되고 응답 지연이 크게 증가하였다.",
        "현재 서비스 규모의 졸업작품 운영/시연 목적에서는 충분한 안정성을 확보한 것으로 판단된다.",
        "향후 더 큰 트래픽을 대비하려면 캐시 적용, DB 쿼리/인덱스 점검, 커넥션 풀 및 서버 리소스 모니터링을 병행하는 것이 적절하다.",
    ]
    for item in conclusions:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item)
        p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "보고서 기재용 요약 문장", 2)
    summary = (
        "운영 서버를 대상으로 k6 기반 트래픽 부하 테스트를 수행한 결과, 동시 가상 사용자 100명 조건에서 실패율 0.00%, "
        "p95 응답 시간 190.65ms로 안정적인 응답 성능을 확인하였다. 150명 조건에서는 실패율 0.00%, p95 613.90ms로 "
        "정상 응답을 유지했으나 응답 지연 증가가 시작되었고, 200명 조건에서는 p95 응답 시간이 1,157.36ms로 증가하여 "
        "처리량 포화 징후가 나타나는 것으로 분석하였다."
    )
    box = doc.add_table(rows=1, cols=1)
    set_table_borders(box, color="CBD5E1")
    set_cell_padding(box, top=160, start=180, bottom=160, end=180)
    set_cell_shading(box.cell(0, 0), "F8FAFC")
    p = box.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, summary, color="111827")


def add_appendix(doc: Document, rows: list[ResultRow]) -> None:
    add_heading(doc, "부록. 결과 파일 및 재현 명령", 1)
    add_body_paragraph(
        doc,
        "정식 측정 결과는 프로젝트의 tools/load-test/results 폴더에 JSON 및 Markdown 파일로 저장되어 있다. "
        "동일 조건 재측정 시 아래 명령 형식을 사용할 수 있다.",
    )
    command = (
        'cd C:\\Users\\leehk\\graduation-project\n'
        '$env:SSLKEYLOGFILE=$null\n'
        '$env:API_BASE_URL="https://newsightkr.com/api"\n'
        '$env:VUS="100"\n'
        '$env:DURATION_SECONDS="120"\n'
        'k6 run tools\\load-test\\newsight-k6.js'
    )
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="CBD5E1")
    set_cell_padding(table, top=140, start=160, bottom=140, end=160)
    set_cell_shading(table.cell(0, 0), "111827")
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(command)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(255, 255, 255)


def build_document() -> Path:
    rows = load_result_rows()
    build_chart(rows)

    doc = Document()
    apply_document_styles(doc)
    add_cover(doc)
    add_summary_box(doc)
    add_environment(doc)
    add_endpoint_table(doc)
    add_result_table(doc, rows)
    add_analysis(doc, rows)
    add_suspected_causes(doc)
    add_conclusion(doc)
    add_appendix(doc, rows)
    add_footer(doc)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT_DOCX)
        return OUTPUT_DOCX
    except PermissionError:
        fallback = OUTPUT_DOCX.with_name("newsight-load-test-report-with-causes-copy.docx")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build_document())
